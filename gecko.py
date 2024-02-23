from requests import Session
from requests.exceptions import HTTPError
from docx import Document
import pandas as pd
import json
import time
import os

def add_cryptos_to_table(table, cryptos, start_index, cryptos_per_page):
    row_cells = None
    for index in range(start_index, start_index + cryptos_per_page):
        if index < len(cryptos):
            crypto = cryptos[index]
            if index % 2 == 0 or index == start_index:
                row_cells = table.add_row().cells
            column_index = index % 2
            row_cells[column_index].text = f"{crypto['Rang']} - {crypto['Nom']} ({crypto['Symbol']})"
    return table

# Pas besoin de charger les variables d'environnement pour CoinGecko
base_url = 'https://api.coingecko.com/api/v3/coins/markets'

# Initialise une liste pour stocker toutes les cryptomonnaies récupérées
all_cryptos = []

try:
    call_count = 0
    page = 1
    limit = 250  # Limite maximale par requête pour CoinGecko
    more_data_available = True

    with Session() as session:
        while more_data_available and call_count < 30:
            params = {
                'vs_currency': 'usd',
                'order': 'market_cap_desc',
                'per_page': limit,
                'page': page
            }

            response = session.get(base_url, params=params)
            if response.status_code == 429:
                print("Rate limit reached. Waiting 60 seconds...")
                time.sleep(60)
                continue
            response.raise_for_status()

            data = response.json()
            if data:
                all_cryptos.extend([{
                    "Rang": str(crypto.get("market_cap_rank")),
                    "Nom": crypto.get("name"),
                    "Symbol": crypto.get("symbol")
                } for crypto in data])
                page += 1
                call_count += 1
                print(f"Page {page} récupérée, total cryptos: {len(all_cryptos)}. Attendre 2 secondes.")
                time.sleep(2)  # Attente pour respecter les limites de taux
            else:
                more_data_available = False

            if call_count == 30:
                print("Limite de 30 appels atteinte. Attente de 30 minutes...")
                time.sleep(1800)  # Attendre 30 minutes
                call_count = 0  # Réinitialiser le compteur d'appels

    with open('cryptomonnaies.json', 'w') as f:
        json.dump(all_cryptos, f)

    df = pd.DataFrame(all_cryptos)
    df.to_excel('cryptomonnaies.xlsx', index=False)

    doc = Document()
    doc.add_heading('Liste des Cryptomonnaies', 0)
    cryptos_per_page = 76
    for i in range(0, len(all_cryptos), cryptos_per_page):
        if i != 0:
            doc.add_page_break()
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        add_cryptos_to_table(table, all_cryptos, i, cryptos_per_page)
    doc.save('cryptomonnaies.docx')

    print(f"{len(all_cryptos)} cryptomonnaies téléchargées et sauvegardées.")

except HTTPError as http_err:
    print(f"Erreur HTTP lors de la connexion à l'API: {http_err}")
except Exception as err:
    print(f"Une erreur s'est produite: {err}")
