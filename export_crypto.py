from requests import Session
from requests.exceptions import HTTPError
from docx import Document
import pandas as pd
import json
from dotenv import load_dotenv
import os

# Fonction pour ajouter des cryptos à une table Word avec deux colonnes
def add_cryptos_to_table(table, cryptos, start_index, cryptos_per_page):
    row_cells = None
    for index in range(start_index, start_index + cryptos_per_page):
        if index < len(cryptos):
            crypto = cryptos[index]
            # Ajouter une nouvelle ligne seulement pour les index pairs ou pour le premier index
            if index % 2 == 0 or index == start_index:
                row_cells = table.add_row().cells
            column_index = index % 2  # 0 pour les index pairs (colonne de gauche), 1 pour les impairs (colonne de droite)
            row_cells[column_index].text = f"{crypto['Rang']} - {crypto['Nom']} ({crypto['Symbol']})"
    return table

# Charger les variables d'environnement
load_dotenv()

# Obtenir les configurations depuis les variables d'environnement
api_key = os.getenv('API_KEY')
base_url = os.getenv('BASE_URL', 'https://pro-api.coinmarketcap.com/v1/cryptocurrency/listings/latest')

if not api_key:
    raise ValueError("La clé API 'API_KEY' n'est pas définie dans les variables d'environnement.")

headers = {
    'Accepts': 'application/json',
    'X-CMC_PRO_API_KEY': api_key,
}

# Initialise une liste pour stocker toutes les cryptomonnaies récupérées
all_cryptos = []

try:
    total_fetched = 0
    limit = 5000  # Limite maximale par requête, ajustez selon les capacités de l'API
    start = 1
    more_data_available = True

    with Session() as session:
        session.headers.update(headers)

        while more_data_available:
            parameters = {
                'start': str(start),
                'limit': str(limit),
                'sort': 'cmc_rank'
            }

            response = session.get(base_url, params=parameters)
            response.raise_for_status()  # Lève une exception pour les réponses d'erreur

            api_response = response.json()
            data = api_response.get('data', [])
            if data:
                all_cryptos.extend([{
                    "Rang": crypto.get("rank"),
                    "Nom": crypto.get("name"),
                    "Symbol": crypto.get("symbol")
                } for crypto in data])
                total_fetched += len(data)
                start += len(data)
            else:
                more_data_available = False

            # Vérifier si on a récupéré moins de données que la limite pour déterminer si c'est la dernière page
            if len(data) < limit:
                more_data_available = False

    # Sauvegarde des données dans un seul fichier JSON
    with open('cryptomonnaies.json', 'w') as f:
        json.dump(all_cryptos, f)

    # Créer un seul fichier Excel avec toutes les cryptomonnaies
    df = pd.DataFrame(all_cryptos)
    df.to_excel('cryptomonnaies.xlsx', index=False)

    # Créer un document Word
    doc = Document()
    doc.add_heading('Liste des Cryptomonnaies', 0)
    cryptos_per_page = 76  # Nombre de cryptomonnaies par page (ajustez ce nombre selon vos besoins)
    for i in range(0, len(all_cryptos), cryptos_per_page):
        if i != 0:
            doc.add_page_break()
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table = add_cryptos_to_table(table, all_cryptos, i, cryptos_per_page)
    doc.save('cryptomonnaies.docx')

    print(f"{total_fetched} cryptomonnaies téléchargées et sauvegardées dans un fichier Excel et un fichier Word.")

except HTTPError as http_err:
    print(f"Erreur HTTP lors de la connexion à l'API: {http_err}")
except Exception as err:
    print(f"Une erreur s'est produite: {err}")
